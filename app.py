import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

# --- 1. 環境構築と設定 ---
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if api_key:
    genai.configure(api_key=api_key)
else:
    st.error("⚠️ APIキーが見つかりません。'.env'ファイルの設定を確認してください。")

# ページ基本設定
st.set_page_config(
    page_title="DocuGen AI | 文章作成アシスタント",
    page_icon="✍️",
    layout="centered"
)

# --- 2. アプリの導入部分（ポートフォリオ用） ---
st.title("📄 DocuGen AI")
st.markdown("### 〜 文章作成の『0→1』を、AIと共に乗り越える 〜")

with st.expander("📌 このアプリについて（開発背景とこだわり）", expanded=True):
    st.write("""
    文章を作成する際、最初の一歩（構成案づくり）に最も時間がかかってしまうことはありませんか？
    このアプリは、そんな**「0→1（ゼロイチ）」の心理的ハードルを下げるため**に開発されました。
    
    **💡 こだわりポイント:**
    1. **論理的な構成:** 提案書には「PREP法」、報告書には「事実と解釈の分離」など、ビジネスフレームワークをAIに徹底させています。
    2. **AIコンサルティング:** 設計書では単なる清書ではなく、技術選定に対するAIからの改善アドバイス機能を搭載しました。
    3. **実用性重視:** ワンクリックコピー機能や、そのまま仕事で使えるWordファイル出力機能を備えています。
    """)

st.divider()

# --- 3. サイドバー：文書形式の選択 ---
st.sidebar.title("🛠 アプリ設定")
doc_type = st.sidebar.selectbox(
    "作成する文書形式を選択",
    ["提案書", "報告書", "設計書"],
    key="doc_type_select"
)

# --- 4. メイン入力フォーム ---
prompt = ""

if doc_type == "提案書":
    st.subheader("💡 提案書の作成")
    st.info("相手の悩みに寄り添い、解決策を論理的に提示する構成を作成します。")
    
    col1, col2 = st.columns(2)
    with col1:
        val1 = st.text_input("1. 相談・依頼内容", placeholder="例：議事録作成の効率化", help="相手からの依頼内容を入力")
    with col2:
        val4 = st.text_input("4. 解決後の理想像", placeholder="例：コア業務への集中時間の増加", help="提案が通った後の成果")
    
    val2 = st.text_area("2. 相手の現状（課題・不満）", placeholder="例：手作業で月20時間かかっている...", height=100)
    val3 = st.text_area("3. 解決策の具体的なアイデア", placeholder="例：AI音声認識ツールの導入...", height=100)
    
    prompt = f"""
    あなたは論理的な提案を得意とするビジネスコンサルタントです。
    相手を動かすための「PREP法」に基づいた提案書を作成してください。
    【相談内容】: {val1}
    【現状の課題】: {val2}
    【提案アイデア】: {val3}
    【理想のゴール】: {val4}
    背景、課題分析、解決策（PREP法）、期待効果、ネクストアクションの順で出力してください。
    """

elif doc_type == "報告書":
    st.subheader("📊 報告書の作成")
    st.info("事実と解釈を分け、リスク分析を含めたプロフェッショナルな報告を行います。")
    
    col1, col2 = st.columns(2)
    with col1:
        val1 = st.text_input("1. 報告の種類", placeholder="例：プロジェクト週報", help="何の報告かを記入")
    with col2:
        val4 = st.text_input("4. 今後の予定", placeholder="例：来週月曜日に修正パッチ適用", help="次に行うべきこと")
    
    val2 = st.text_area("2. 実施事項（事実）", placeholder="例：API連携テストの実施...", height=100)
    val3 = st.text_area("3. 結果・判明したこと（解釈）", placeholder="例：テストは成功したが、負荷時に遅延あり...", height=100)
    
    prompt = f"""
    あなたは優秀なプロジェクトマネージャーです。
    【種類】: {val1}
    【実施事項】: {val2}
    【結果（事実と解釈）】: {val3}
    【予定】: {val4}
    上記を整理し、冒頭に「要約ステータス」を置き、結果に基づいた「潜在的リスクと対策」をAIの視点で分析して含めてください。
    """

elif doc_type == "設計書":
    st.subheader("🏗 設計書の作成")
    st.info("シニアエンジニアの視点で、設計案と技術選定のアドバイスを出力します。")
    
    col1, col2 = st.columns(2)
    with col1:
        val1 = st.text_input("1. システム・機能名", placeholder="例：在庫管理アプリ")
    with col2:
        val3 = st.text_input("3. 使いたい技術（AIが添削します）", placeholder="例：HTML, SQLite")
    
    val2 = st.text_area("2. 作成の目的", placeholder="例：手書き業務のデジタル化...", height=100)
    val4 = st.text_area("4. 主要な機能", placeholder="例：バーコードスキャン、通知機能...", height=100)
    
    prompt = f"""
    あなたはシニアシステムアーキテクトです。以下の要件に基づき設計草案を作成してください。
    【システム名】: {val1}
    【目的】: {val2}
    【技術構成】: {val3}
    【機能】: {val4}
    出力には必ず「技術選定へのアドバイス」セクションを設け、ユーザーの選んだ技術({val3})が最適か評価し、より良い代替案があれば理由とともに提案してください。
    """

# --- 5. 生成実行と表示 ---
st.divider()
if st.button(f"✨ {doc_type}を生成する", type="primary", key="generate_button", use_container_width=True):
    if not all([val1, val2, val3, val4]):
        st.warning("⚠️ すべての項目を入力してください。")
    else:
        with st.spinner("AIが構成を練っています..."):
            try:
                # Gemini 2.5 Flashを使用（2026年最新環境）
                model = genai.GenerativeModel("models/gemini-2.5-flash")
                response = model.generate_content(prompt)
                
                st.success(f"✅ {doc_type}の草案が完成しました！")
                
                # 結果表示エリア
                with st.container(border=True):
                    st.subheader("📋 生成結果")
                    st.code(response.text, language="markdown", wrap_lines=True)
                    
                    with st.expander("👁 リッチテキスト形式でプレビュー"):
                        st.markdown(response.text)

                # Wordファイル作成
                doc = Document()
                doc.add_heading(f'{doc_type} 草案', 0)
                doc.add_paragraph(response.text)
                bio = BytesIO()
                doc.save(bio)

                # ダウンロードボタン
                st.download_button(
                    label="📄 Wordファイルとして保存",
                    data=bio.getvalue(),
                    file_name=f"{doc_type}_草案.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_button",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"エラーが発生しました: {e}")

# フッター
st.caption("© 2026 DocuGen AI - Powered by Gemini 2.5 Flash")